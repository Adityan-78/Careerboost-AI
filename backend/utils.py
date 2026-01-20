import fitz  # PyMuPDF
from docx import Document
import logging
from io import BytesIO
import re

logger = logging.getLogger(__name__)


def clean_extracted_text(text: str) -> str:
    """
    Clean and normalize extracted text to ensure consistency.
    
    Args:
        text: Raw extracted text
    
    Returns:
        Cleaned text
    """
    # Remove null bytes and control characters except newlines and tabs
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    
    # Normalize bullet points and special characters
    text = text.replace('•', '* ')
    text = text.replace('·', '* ')
    text = text.replace('◦', '* ')
    text = text.replace('▪', '* ')
    text = text.replace('‣', '* ')
    text = text.replace('⚫', '* ')
    
    # Normalize whitespace - replace multiple spaces with single space
    text = re.sub(r' +', ' ', text)
    
    # Normalize line breaks - remove excessive blank lines
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    return text.strip()


def parse_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file using PyMuPDF with improved formatting.
    Uses multiple extraction methods for better reliability.
    
    Args:
        file_content: PDF file content as bytes
    
    Returns:
        Extracted text as string
    """
    pdf_document = None
    try:
        # Open PDF from bytes
        pdf_document = fitz.open(stream=file_content, filetype="pdf")
        
        text_content = []
        page_count = pdf_document.page_count
        
        # Extract text from each page
        for page_num in range(page_count):
            page = pdf_document[page_num]
            
            # Method 1: Try standard text extraction
            page_text = page.get_text("text")
            
            # Method 2: If standard extraction yields poor results, try blocks
            if len(page_text.strip()) < 50:  # Likely poor extraction
                blocks = page.get_text("blocks")
                block_texts = []
                for block in blocks:
                    # block[4] contains the text
                    if len(block) > 4 and block[4].strip():
                        block_texts.append(block[4].strip())
                page_text = '\n'.join(block_texts)
            
            # Method 3: If still poor, try dict extraction
            if len(page_text.strip()) < 50:
                text_dict = page.get_text("dict")
                dict_texts = []
                for block in text_dict.get("blocks", []):
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line.get("spans", []):
                                span_text = span.get("text", "").strip()
                                if span_text:
                                    dict_texts.append(span_text)
                page_text = ' '.join(dict_texts)
            
            # Clean and add to content
            if page_text.strip():
                cleaned_text = clean_extracted_text(page_text)
                if cleaned_text:
                    text_content.append(cleaned_text)
        
        # Join all pages BEFORE closing the document
        full_text = "\n\n".join(text_content)
        
        # Now it's safe to close
        pdf_document.close()
        
        if not full_text.strip():
            raise ValueError("No text content could be extracted from PDF. The PDF may be image-based or encrypted.")
        
        # Final cleaning
        full_text = clean_extracted_text(full_text)
        
        logger.info(f"Successfully extracted {len(full_text)} characters from PDF ({page_count} pages)")
        return full_text
    
    except Exception as e:
        logger.error(f"PDF parsing error: {str(e)}")
        # Make sure document is closed even on error
        if pdf_document is not None:
            try:
                pdf_document.close()
            except:
                pass
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def parse_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file_content: DOCX file content as bytes
    
    Returns:
        Extracted text as string
    """
    try:
        # Open DOCX from bytes
        doc = Document(BytesIO(file_content))
        
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        # Join all content with newlines
        full_text = "\n".join(text_content)
        
        if not full_text.strip():
            raise ValueError("No text content found in DOCX")
        
        # Clean the text
        full_text = clean_extracted_text(full_text)
        
        logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
        return full_text
    
    except Exception as e:
        logger.error(f"DOCX parsing error: {str(e)}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def parse_file(file_content: bytes, file_extension: str) -> str:
    """
    Parse resume file based on extension.
    
    Args:
        file_content: File content as bytes
        file_extension: File extension (e.g., '.pdf', '.docx')
    
    Returns:
        Extracted text as string
    """
    file_extension = file_extension.lower()
    
    if file_extension == ".pdf":
        return parse_pdf(file_content)
    elif file_extension in [".docx", ".doc"]:
        return parse_docx(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")